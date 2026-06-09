import re
from pathlib import Path
from typing import Any, Dict, List

import fitz
from docx import Document


_CHECKBOX_RE = re.compile(r"^[\s□■☑☒]+")
_CHECKBOX_ANY_RE = re.compile(r"[□■☑☒]")
_MIN_INDENT_DELTA_PT = 6.0


def extract_checkboxes(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Estrae le caselle da spuntare (checkbox) dai blocchi.
    Considera checkbox solo i blocchi con testo che inizia con simboli □■☑☒.
    Accorpa le righe successive che risultano indentate (continuazione multi-riga).
    """
    items: List[Dict[str, Any]] = []
    if not blocks:
        return items

    for i, block in enumerate(blocks):
        text = str(block.get("text", "")).strip()
        if not text:
            continue
        if _CHECKBOX_RE.search(text) is None:
            continue

        page = block.get("page", 1)
        bbox = block.get("bbox", [0, 0, 0, 0])
        start_x = float(bbox[0]) if bbox else 0.0
        lines = [text]

        # Accorpa righe successive indentate sulla stessa pagina
        prev = block
        for next_block in blocks[i + 1:]:
            if next_block.get("page", 1) != page:
                break
            next_text = str(next_block.get("text", "")).strip()
            if not next_text:
                prev = next_block
                continue
            # Stop se inizia un'altra checkbox
            if _CHECKBOX_RE.search(next_text):
                break

            next_bbox = next_block.get("bbox", [0, 0, 0, 0])
            next_x = float(next_bbox[0]) if next_bbox else 0.0
            if next_x >= start_x + _MIN_INDENT_DELTA_PT:
                lines.append(next_text)
                prev = next_block
                continue
            break

        label = _clean_checkbox_label(text)
        items.append(
            {
                "label": label,
                "text": " ".join(lines).strip(),
                "lines": lines,
                "page": page,
                "bbox": bbox,
                "marker_bbox": _marker_bbox_from_block_bbox(bbox, text),
            }
        )

    return items


def _clean_checkbox_label(text: str) -> str:
    text = _CHECKBOX_RE.sub("", text).strip()
    text = text.strip(" :,;./\\()")
    return text.strip()


def extract_checkboxes_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Estrae checkbox da un DOCX usando python-docx.
    Riconosce simboli checkbox e controlli checkBox nel XML.
    Accorpa paragrafi successivi indentati come continuazione.
    """
    doc = Document(docx_path)
    items: List[Dict[str, Any]] = []

    # Paragraphs
    items.extend(_extract_from_paragraphs(list(doc.paragraphs)))

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_items = _extract_from_paragraphs(list(cell.paragraphs), allow_indent=False)
                items.extend(cell_items)

    # Dedupe by (label, text)
    seen = set()
    unique: List[Dict[str, Any]] = []
    for item in items:
        key = (item.get("label"), item.get("text"))
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)

    return unique


def _extract_from_paragraphs(paragraphs, allow_indent: bool = True) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if not _paragraph_has_checkbox(p):
            i += 1
            continue

        text = _paragraph_text_with_checkbox(p)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        start_indent = _indent_pt(p)

        j = i + 1
        while j < len(paragraphs):
            p2 = paragraphs[j]
            if _paragraph_has_checkbox(p2):
                break
            if not p2.text or not p2.text.strip():
                j += 1
                continue
            if allow_indent and _indent_pt(p2) >= start_indent + _MIN_INDENT_DELTA_PT:
                lines.extend([line.strip() for line in p2.text.splitlines() if line.strip()])
                j += 1
                continue
            if not allow_indent:
                lines.extend([line.strip() for line in p2.text.splitlines() if line.strip()])
                j += 1
                continue
            break

        full_text = " ".join(lines).strip()
        label = _clean_checkbox_label(lines[0]) if lines else ""
        items.append(
            {
                "label": label,
                "text": full_text,
                "lines": lines,
                "page": 1,
                "bbox": None,
                "marker_bbox": None,
            }
        )
        i = j

    return items


def _paragraph_has_checkbox(paragraph) -> bool:
    # Testo diretto
    if _CHECKBOX_ANY_RE.search(paragraph.text or ""):
        return True
    # Simboli in XML (w:sym o controllo checkbox)
    xml = paragraph._p.xml
    if "<w:sym" in xml:
        return True
    if "<w:checkBox" in xml or "<w14:checkbox" in xml:
        return True
    return False


def _paragraph_text_with_checkbox(paragraph) -> str:
    text = paragraph.text or ""
    if _CHECKBOX_ANY_RE.search(text):
        return text
    # Se il simbolo non è nel testo, premetti un placeholder checkbox.
    return f"□ {text}".strip()


def _indent_pt(paragraph) -> float:
    indent = paragraph.paragraph_format.left_indent
    if indent is None:
        return 0.0
    return float(indent.pt)


def _marker_bbox_from_block_bbox(bbox, text: str):
    if not bbox or len(bbox) != 4:
        return None
    x, y, w, h = [float(value) for value in bbox]
    stripped = (text or "").lstrip()
    if not stripped:
        return None
    if stripped[0] not in {"□", "■", "☑", "☒"}:
        return None
    size = max(6.0, min(h, 14.0))
    return [round(x, 2), round(y, 2), round(size, 2), round(h, 2)]

def extract_checkboxes_from_pdf(pdf_path, original_path=None, output_dir=None, raster_mode="off") -> List[Dict[str, Any]]:
    doc = fitz.open(str(pdf_path))
    items: List[Dict[str, Any]] = []

    for page in doc:
        for drawing in page.get_drawings():
            rect = drawing.get("rect")
            if not rect:
                continue

            w = float(rect.width)
            h = float(rect.height)
            if w < 4 or h < 4 or w > 24 or h > 24:
                continue
            if abs(w - h) / max(w, h) > 0.38:
                continue

            text_right = ""
            words = []
            cy = (rect.y0 + rect.y1) / 2.0
            for raw_word in page.get_text("words"):
                word_rect = fitz.Rect(raw_word[:4])
                word_cy = (word_rect.y0 + word_rect.y1) / 2.0
                if abs(word_cy - cy) <= max(5.0, h * 0.85) and word_rect.x0 > rect.x1:
                    words.append((word_rect.x0, raw_word[4]))
            words.sort(key=lambda x: x[0])
            text_right = " ".join(word for _, word in words).strip()

            bbox = [round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2)]
            items.append(
                {
                    "label": text_right,
                    "text": text_right,
                    "lines": [text_right] if text_right else [],
                    "page": page.number + 1,
                    "bbox": bbox,
                    "marker_bbox": bbox,
                    "checkbox_bbox": bbox,
                }
            )

    doc.close()
    return items
