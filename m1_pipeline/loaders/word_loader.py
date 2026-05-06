import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt


# Altezza riga approssimata in punti per stimare le coordinate y
_LINE_HEIGHT_PT = 14


def convert_doc_to_docx(doc_path: str) -> str:
    """
    Converte un file .doc in .docx usando Microsoft Word via COM automation.
    Richiede Microsoft Word installato su Windows.

    Args:
        doc_path: Percorso al file .doc.

    Returns:
        Percorso al file .docx convertito (in una cartella temporanea).

    Raises:
        RuntimeError: se Microsoft Word non è disponibile o la conversione fallisce.
    """
    try:
        import win32com.client as win32
    except ImportError:
        raise RuntimeError(
            "pywin32 non installato. Esegui: pip install pywin32"
        )

    doc_path = str(Path(doc_path).resolve())
    tmp_dir = tempfile.mkdtemp()
    docx_path = str(Path(tmp_dir) / (Path(doc_path).stem + ".docx"))

    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0          # wdAlertsNone
        word.Options.ConfirmConversions = False

        pv = word.Application.FileDialog  
        try:
            word.Options.DoNotPromptForConvert = True
        except Exception:
            pass
        try:
            word.Application.ProtectedViewWindows  
            word.Options.ProtectedView = 0
        except Exception:
            pass

        doc = word.Documents.Open(
            doc_path,   # FileName
            False,      # ConfirmConversions
            False,      # ReadOnly
            False,      # AddToRecentFiles
        )

        if doc.__class__.__name__ != "Document":
            try:
                pv_win = word.Application.ProtectedViewWindows(1)
                doc = pv_win.Edit()
            except Exception:
                pass

        try:
            doc.SaveAs2(docx_path, FileFormat=12)
        except AttributeError:
            doc.SaveAs(docx_path, FileFormat=12)

        doc.Close(0)  # wdDoNotSaveChanges
        return docx_path
    except Exception as e:
        raise RuntimeError(
            f"Conversione .doc → .docx fallita. Assicurati che Microsoft Word sia installato.\n"
            f"Dettaglio: {e}"
        )
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def convert_word_to_pdf(word_path: str) -> str:
    try:
        import win32com.client as win32
    except ImportError:
        raise RuntimeError(
            "pywin32 non installato. Esegui: pip install pywin32"
        )

    word_path = str(Path(word_path).resolve())
    tmp_dir = tempfile.mkdtemp()
    pdf_path = str(Path(tmp_dir) / (Path(word_path).stem + ".pdf"))

    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        word.Options.ConfirmConversions = False

        doc = word.Documents.Open(
            word_path,
            False,
            True,
            False,
        )
        try:
            doc.SaveAs2(pdf_path, FileFormat=17)
        except AttributeError:
            doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close(0)
        return pdf_path
    except Exception as e:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise RuntimeError(
            f"Conversione Word → PDF fallita. Assicurati che Microsoft Word sia installato.\n"
            f"Dettaglio: {e}"
        )
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def load_word(file_path: str) -> List[Dict[str, Any]]:
    """
    Estrae testo strutturato da un file Word (.docx o .doc).
    I file .doc vengono convertiti automaticamente in .docx prima dell'elaborazione.

    Args:
        file_path: Percorso al file .docx o .doc.

    Returns:
        Lista di blocchi nel formato standard:
        {"text": str, "bbox": [x, y, w, h], "page": int, "confidence": float, "style": str}
    """
    suffix = Path(file_path).suffix.lower()
    cleanup_dirs = []

    if suffix == ".doc":
        print(f"[INFO] File .doc rilevato — conversione in .docx in corso...")
        docx_path = convert_doc_to_docx(file_path)
        cleanup_dirs.append(str(Path(docx_path).parent))
        file_path = docx_path
        print(f"[INFO] Conversione completata: {docx_path}")

    try:
        pdf_path = convert_word_to_pdf(file_path)
        cleanup_dirs.append(str(Path(pdf_path).parent))
        print(f"[INFO] Conversione Word -> PDF completata: {pdf_path}")
        from loaders.pdf_loader import load_pdf
        blocks = load_pdf(pdf_path)
    except Exception:
        blocks = _extract_blocks(file_path)
    finally:
        for directory in cleanup_dirs:
            shutil.rmtree(directory, ignore_errors=True)

    return blocks


def _extract_blocks(file_path: str) -> List[Dict[str, Any]]:
    """Estrae blocchi da un file .docx."""
    doc = Document(file_path)
    blocks: List[Dict[str, Any]] = []
    y_cursor = 0.0

    # --- Paragrafi ---
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            y_cursor += _LINE_HEIGHT_PT
            continue

        font_size = _get_font_size(paragraph)
        line_h = font_size if font_size else _LINE_HEIGHT_PT

        blocks.append(
            {
                "text": text,
                "bbox": [0.0, round(y_cursor, 2), 500.0, round(line_h, 2)],
                "page": 1,
                "confidence": 1.0,
                "style": paragraph.style.name,
                "source": "word",
            }
        )
        y_cursor += line_h + 2

    # --- Tabelle ---
    for table in doc.tables:
        for row in table.rows:
            x_cursor = 0.0
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    blocks.append(
                        {
                            "text": text,
                            "bbox": [round(x_cursor, 2), round(y_cursor, 2), 120.0, _LINE_HEIGHT_PT],
                            "page": 1,
                            "confidence": 1.0,
                            "style": "table_cell",
                            "source": "word",
                        }
                    )
                x_cursor += 120.0
            y_cursor += _LINE_HEIGHT_PT + 2

    return blocks


def _get_font_size(paragraph) -> float:
    """Restituisce la dimensione font del primo run del paragrafo, o None."""
    for run in paragraph.runs:
        if run.font.size:
            return run.font.size.pt
    return None
